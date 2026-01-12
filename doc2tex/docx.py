# doc2tex - LaTeX to Word conversion logic
# To be honest, this way is harder because LaTeX is basically a programming language.
# I'm using regex to find the main parts, it's not perfect but it handles 
# normal documents pretty well.

import os
import re
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .options import ConversionOptions
from .utils import unescape_latex, logger, ensure_directory
from .errors import ConversionError


class DocxGenerator:
    """
    My class for generating Word docs from LaTeX.
    It's basically a simple parser that looks for commands like \section or \textbf.
    """
    
    def __init__(self, options: ConversionOptions):
        self.options = options
        self.word_doc = None
        
    def convert(self, tex_file: str, docx_file: str) -> str:
        # Tries to turn your .tex into a .docx
        try:
            logger.info(f"Trying to read {tex_file}...")
            
            # Read all the text
            with open(tex_file, 'r', encoding=self.options.output_encoding) as f:
                raw_tex = f.read()
            
            # Create a blank Word document
            self.word_doc = Document()
            
            # Set the font to something standard (students love Times New Roman)
            self._apply_student_styles()
            
            # This is where the magic (or mess) happens
            self._parse_and_build(raw_tex)
            
            # Save the result
            self.word_doc.save(docx_file)
            logger.info(f"Nice! Saved the word doc to {docx_file}")
            
            return docx_file
            
        except Exception as err:
            logger.error(f"LaTeX parsing failed: {err}")
            raise ConversionError(f"Something went wrong reading the LaTeX file: {err}")

    def _apply_student_styles(self) -> None:
        # Setup the document styles to look like a standard report
        style = self.word_doc.styles['Normal']
        f = style.font
        f.name = 'Times New Roman'
        
        # Pull font size from options (usually 12pt)
        try:
             # Just stripping 'pt' if it's there
             sz = int(self.options.font_size.value.replace('pt', ''))
             f.size = Pt(sz)
        except:
             f.size = Pt(12) # fallback

    def _parse_and_build(self, content: str) -> None:
        # We only really care about stuff inside \begin{document}
        # If we can't find it, we just take everything
        pattern = r'\\begin\{document\}(.*?)\\end\{document\}'
        match = re.search(pattern, content, re.DOTALL)
        
        if match:
            doc_body = match.group(1).strip()
        else:
            # Maybe it's just a snippet?
            doc_body = content.strip()
            
        # I split the body into blocks by double newlines
        # This usually means separate paragraphs or sections in LaTeX
        blocks = re.split(r'\n\s*\n', doc_body)
        
        for bk in blocks:
            bk = bk.strip()
            if not bk:
                continue
                
            # Figure out what this block is
            if bk.startswith('\\section'):
                self._add_heading(bk, 1)
            elif bk.startswith('\\subsection'):
                self._add_heading(bk, 2)
            elif bk.startswith('\\subsubsection'):
                self._add_heading(bk, 3)
            elif '\\begin{table}' in bk:
                self._add_table(bk)
            elif '\\begin{figure}' in bk:
                self._add_image(bk)
            elif '\\begin{itemize}' in bk or '\\begin{enumerate}' in bk:
                self._add_list(bk)
            elif '\\begin{center}' in bk:
                self._add_centered(bk)
            else:
                # If it's none of the above, it's probably just a normal paragraph
                self._add_paragraph(bk)

    def _add_heading(self, block: str, level: int) -> None:
        # Extract text from \section{...} or \subsection{...}
        m = re.search(r'\\(?:sub)*section\{([^}]+)\}', block)
        if m:
            title = unescape_latex(m.group(1))
            self.word_doc.add_heading(title, level=level)

    def _add_paragraph(self, block: str) -> None:
        # Handles text along with inline styles like bold/italic
        p = self.word_doc.add_paragraph()
        self._apply_inline(block, p)

    def _apply_inline(self, text: str, para_obj) -> None:
        # This is my favorite part: a simple inline 'parser'
        # It looks for formatting tags and adds them as 'runs'
        
        # These are the things we support right now
        patterns = [
            (r'\\textbf\{([^}]+)\}', 'bold'),
            (r'\\textit\{([^}]+)\}', 'italic'),
            (r'\\underline\{([^}]+)\}', 'underline'),
            (r'\$([^$]+)\$', 'math'), # Simple inline math between $$
        ]
        
        idx = 0
        while idx < len(text):
            found_m = None
            found_type = None
            
            # Check all patterns to see which one comes next in the string
            for pat, t in patterns:
                m = re.search(pat, text[idx:])
                if m:
                    if not found_m or m.start() < found_m.start():
                        found_m = m
                        found_type = t
            
            if not found_m:
                # No more formatting tags, just add the rest and finish
                rest = unescape_latex(text[idx:])
                if rest:
                    para_obj.add_run(rest)
                break
                
            # Add the text BEFORE the formatting tag
            pre = unescape_latex(text[idx : idx + found_m.start()])
            if pre:
                para_obj.add_run(pre)
            
            # Handle the actual formatted text
            content = unescape_latex(found_m.group(1))
            run = para_obj.add_run(content)
            
            if found_type == 'bold':
                run.bold = True
            elif found_type == 'italic':
                run.italic = True
            elif found_type == 'underline':
                run.underline = True
            elif found_type == 'math':
                # For math, we just make it italic for now so it looks different
                run.italic = True
                
            # Move the index past this match
            idx += found_m.end()

    def _add_table(self, block: str) -> None:
        # Tries to rebuild a table from tabular
        tab_m = re.search(r'\\begin\{tabular\}\{[^}]+\}(.*?)\\end\{tabular\}', block, re.DOTALL)
        if not tab_m:
             return
             
        rows_text = tab_m.group(1).strip()
        
        # Split rows by \\ (the LaTeX row separator)
        # Note: I'm skipping common lines like \hline or \midrule
        lines = [r.strip() for r in rows_text.split('\\\\') if r.strip()]
        lines = [r for r in lines if not r.startswith('\\')]
        
        if not lines:
             return
             
        # Guess how many columns based on the first row
        first_row = lines[0].split('&')
        num_c = len(first_row)
        
        # Add the table to Word
        t = self.word_doc.add_table(rows=len(lines), cols=num_c)
        t.style = 'Table Grid'
        
        for r_idx, r_text in enumerate(lines):
            cells = [unescape_latex(c.strip()) for c in r_text.split('&')]
            for c_idx, val in enumerate(cells):
                if c_idx < num_c:
                     t.rows[r_idx].cells[c_idx].text = val

    def _add_image(self, block: str) -> None:
        # Looks for \includegraphics and adds the picture to Word
        m = re.search(r'\\includegraphics(?:\[[^\]]+\])?\{([^}]+)\}', block)
        if m:
            img_path = m.group(1)
            # We check if the file actually exists
            if os.path.exists(img_path):
                 try:
                      self.word_doc.add_picture(img_path, width=Inches(4))
                 except:
                      self.word_doc.add_paragraph(f"[Image found but error loading: {img_path}]")
            else:
                 self.word_doc.add_paragraph(f"[Image file not found: {img_path}]")

    def _add_list(self, block: str) -> None:
        # Reconstructs bullet/numbered lists
        is_num = '\\begin{enumerate}' in block
        # Find every \item text
        items = re.findall(r'\\item\s+(.*?)(?=\\item|\n|\\end)', block, re.DOTALL)
        
        for it in items:
            style = 'List Number' if is_num else 'List Bullet'
            p = self.word_doc.add_paragraph(style=style)
            self._apply_inline(it.strip(), p)

    def _add_centered(self, block: str) -> None:
        m = re.search(r'\\begin\{center\}(.*?)\\end\{center\}', block, re.DOTALL)
        if m:
             txt = m.group(1).strip()
             p = self.word_doc.add_paragraph()
             p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
             self._apply_inline(txt, p)
